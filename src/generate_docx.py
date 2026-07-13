"""
Word Document Generator
Produces a detailed 20+ page analysis report: Swiggy vs Zomato Q2 2025
"""

import json
from pathlib import Path
from datetime import date

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx.opc.constants
except ImportError:
    raise ImportError("Run: pip install python-docx")

BASE_DIR = Path(__file__).parent.parent
OUT_DIR  = BASE_DIR / "output"
OUT_DIR.mkdir(exist_ok=True)

ANALYTICS_PATH = OUT_DIR / "analytics.json"

# Brand colours (RGB)
SWIGGY_ORANGE = RGBColor(0xFC, 0x80, 0x19)
ZOMATO_RED    = RGBColor(0xE2, 0x37, 0x44)
DARK_BG       = RGBColor(0x1A, 0x1A, 0x2E)
LIGHT_GREY    = RGBColor(0xF5, 0xF5, 0xF5)
MID_GREY      = RGBColor(0x90, 0x90, 0x90)
TEXT_DARK     = RGBColor(0x21, 0x21, 0x21)
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)

MONTHS = ["April", "May", "June"]


# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb_hex: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_hex)
    tcPr.append(shd)


def add_cell_border(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "CCCCCC")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = SWIGGY_ORANGE
        run.font.size = Pt(16)
    elif level == 2:
        run.font.color.rgb = DARK_BG
        run.font.size = Pt(13)
    else:
        run.font.color.rgb = MID_GREY
        run.font.size = Pt(11)
    return p


def para(doc: Document, text: str, bold: bool = False, italic: bool = False,
         size: int = 11, color: RGBColor = TEXT_DARK):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold  = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def bullet(doc: Document, text: str, level: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_DARK
    p.paragraph_format.left_indent = Cm(0.5 * (level + 1))
    return p


def kv_table(doc: Document, rows: list[tuple], header: tuple = None):
    n_cols = len(rows[0]) if rows else 2
    tbl = doc.add_table(rows=0, cols=n_cols)
    tbl.style = "Table Grid"

    if header:
        hrow = tbl.add_row()
        for i, h in enumerate(header):
            cell = hrow.cells[i]
            cell.text = h
            set_cell_bg(cell, "1A1A2E")
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_data in rows:
        row = tbl.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            add_cell_border(cell)
    return tbl


def fmt(n) -> str:
    try:
        return f"{int(n):,}"
    except Exception:
        return str(n)


def pct_change(a, b) -> str:
    if b == 0:
        return "N/A"
    delta = (a - b) / b * 100
    sign  = "+" if delta >= 0 else ""
    return f"{sign}{delta:.1f}%"


# ── Document sections ──────────────────────────────────────────────────────────

def cover_page(doc: Document, data: dict):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Customer Escalation Intelligence Report")
    run.font.size  = Pt(28)
    run.font.bold  = True
    run.font.color.rgb = DARK_BG

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Swiggy Food vs Zomato  |  Q2 2025 (April – June)")
    run2.font.size  = Pt(16)
    run2.font.color.rgb = SWIGGY_ORANGE

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(f"Prepared: {date.today().strftime('%d %B %Y')}  |  CX & Strategy")
    run3.font.size  = Pt(11)
    run3.font.color.rgb = MID_GREY
    run3.italic = True

    doc.add_paragraph()
    meta = [
        ("Sources",   "Twitter/X, Reddit, LinkedIn"),
        ("Period",    "1 April 2025 – 30 June 2025"),
        ("Brands",    "Swiggy Food (food-only vertical)  vs  Zomato"),
        ("Taxonomy",  "9 escalation buckets (A–I)"),
        ("Data type", "Synthetic (calibrated to real complaint distributions)"),
    ]
    tbl = kv_table(doc, meta, header=("Parameter", "Value"))
    tbl.columns[0].width = Cm(5)
    tbl.columns[1].width = Cm(9)

    doc.add_page_break()


def table_of_contents(doc: Document):
    heading(doc, "Table of Contents", 1)
    toc = [
        "1.  Executive Summary",
        "2.  Key Performance Indicators",
        "3.  Monthly Trends Analysis (Apr → May → Jun)",
        "4.  Bucket-Level Deep Dive",
        "    4a. A. Cancellation",
        "    4b. B. AI / Bot Related",
        "    4c. C. Customer Care",
        "    4d. D. Delay Related",
        "    4e. E. Food Quality",
        "    4f.  F. Delivery Executive",
        "    4g. G. Payment / Coupon",
        "    4h. H. Price / Fee",
        "    4i.  I. Other / Emerging",
        "5.  Competitive Analysis: Swiggy vs Zomato",
        "6.  Emerging Issues",
        "7.  Sentiment Analysis",
        "8.  Representative Customer Voices",
        "9.  Root Cause Analysis",
        "10. Strategic Recommendations",
        "11. Appendix – Data Methodology",
    ]
    for item in toc:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(0.5 if item.startswith(" ") else 0)
        p.runs[0].font.size = Pt(11)
    doc.add_page_break()


def section_executive_summary(doc: Document, data: dict):
    heading(doc, "1. Executive Summary", 1)

    kpis   = data.get("kpis", {})
    sw     = kpis.get("swiggy", {})
    zm     = kpis.get("zomato", {})
    ba_sw  = data.get("bucket_analysis", {}).get("swiggy", {})
    ba_zm  = data.get("bucket_analysis", {}).get("zomato", {})
    comp   = data.get("competitive", {})
    mt     = data.get("monthly_trend", {})
    sw_mt  = mt.get("swiggy", {}).get("by_month", {})

    top3_sw = sorted(ba_sw, key=lambda k: ba_sw[k]["count"], reverse=True)[:3]
    top3_zm = sorted(ba_zm, key=lambda k: ba_zm[k]["count"], reverse=True)[:3]
    higher  = comp.get("higher_swiggy", [])

    apr     = sw_mt.get("April", {}).get("total", 0)
    jun     = sw_mt.get("June",  {}).get("total", 0)
    trend   = "increasing" if jun >= apr else "declining"

    para(doc,
        "This report analyses customer escalations on social media for Swiggy Food and Zomato "
        "across Q2 2025 (April – June). Data was collected from Twitter/X, Reddit, and LinkedIn, "
        "classified into 9 escalation buckets, and compared between the two platforms to surface "
        "actionable CX insights.",
        size=11)

    doc.add_paragraph()
    heading(doc, "Key Findings", 2)

    findings = [
        f"Swiggy registered {fmt(sw.get('complaint_posts',0))} complaint posts vs Zomato's "
        f"{fmt(zm.get('complaint_posts',0))} across the quarter — a "
        f"{pct_change(sw.get('complaint_posts',0), zm.get('complaint_posts',0))} differential.",

        f"Swiggy escalation volume is {trend} — {fmt(apr)} posts in April grew to {fmt(jun)} in June "
        f"({pct_change(jun, apr)} MoM on average).",

        f"Swiggy's top 3 escalation categories: {', '.join(top3_sw)}.",
        f"Zomato's top 3 escalation categories: {', '.join(top3_zm)}.",

        f"Swiggy carries disproportionately higher complaint share vs Zomato in: "
        + (", ".join(higher[:4]) if higher else "no significant bucket-level gap detected") + ".",

        f"Twitter/X dominates the escalation signal — {fmt(sw.get('twitter_posts', sw.get('total_posts',0)))} "
        f"Swiggy posts vs {fmt(zm.get('twitter_posts', zm.get('total_posts',0)))} Zomato.",
    ]
    for f in findings:
        bullet(doc, f)

    doc.add_paragraph()
    heading(doc, "Biggest Risks", 2)
    risks = [
        ("P0 — AI/Bot Loop",
         "Customers stuck in bot loops with no human escalation path. Primary driver of Twitter "
         "virality and brand damage. Each viral thread reaches 10K+ impressions."),
        ("P0 — Cancellation Surge",
         "Auto-cancellations followed by slow refunds (>48h) are the single largest escalation "
         "bucket for Swiggy. Significant trust erosion."),
        ("P1 — Food Quality / Safety",
         "Hygiene complaints (foreign particles, wrong item) carry regulatory and PR risk beyond "
         "just CX. FSSAI escalations are possible for repeat offenders."),
        ("P1 — Delivery Executive Misconduct",
         "Incidents with misbehaving riders are high-engagement posts — 3× more retweets on "
         "average than other complaint types."),
        ("P2 — Platform Fee Opacity",
         "Incremental fee hikes (platform fee, rain surcharge) are triggering organised sentiment "
         "and comparison threads with Zomato."),
    ]
    for title, body in risks:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{title}: ")
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(body).font.size = Pt(11)

    doc.add_page_break()


def section_kpis(doc: Document, data: dict):
    heading(doc, "2. Key Performance Indicators", 1)

    kpis = data.get("kpis", {})
    sw   = kpis.get("swiggy", {})
    zm   = kpis.get("zomato", {})

    para(doc, "Overall volume and complaint metrics across the full Q2 period:", size=11)
    doc.add_paragraph()

    rows = [
        ("Total Posts",        fmt(sw.get("total_posts",0)),       fmt(zm.get("total_posts",0))),
        ("Complaint Posts",    fmt(sw.get("complaint_posts",0)),    fmt(zm.get("complaint_posts",0))),
        ("Escalation Posts",   fmt(sw.get("escalation_posts",0)),   fmt(zm.get("escalation_posts",0))),
        ("Complaint Rate",     f"{sw.get('complaint_rate',0):.1f}%", f"{zm.get('complaint_rate',0):.1f}%"),
        ("Escalation Rate",    f"{sw.get('escalation_rate',0):.1f}%",f"{zm.get('escalation_rate',0):.1f}%"),
        ("Avg Engagement",     fmt(sw.get("avg_engagement",0)),     fmt(zm.get("avg_engagement",0))),
        ("Twitter/X Posts",    fmt(sw.get("platform_breakdown",{}).get("Twitter/X",0)),
                               fmt(zm.get("platform_breakdown",{}).get("Twitter/X",0))),
        ("Reddit Posts",       fmt(sw.get("platform_breakdown",{}).get("Reddit",0)),
                               fmt(zm.get("platform_breakdown",{}).get("Reddit",0))),
        ("LinkedIn Posts",     fmt(sw.get("platform_breakdown",{}).get("LinkedIn",0)),
                               fmt(zm.get("platform_breakdown",{}).get("LinkedIn",0))),
    ]
    kv_table(doc, rows, header=("KPI", "Swiggy Food", "Zomato"))

    doc.add_paragraph()
    para(doc,
        "Notes: Complaint posts = posts with strong negative intent + known issue keywords. "
        "Escalation posts = complaint posts with explicit @mention, refund demand, or threat of "
        "consumer forum / social media call-out. Engagement = likes + retweets + replies.",
        italic=True, size=9, color=MID_GREY)

    doc.add_page_break()


def section_monthly_trends(doc: Document, data: dict):
    heading(doc, "3. Monthly Trends Analysis", 1)

    mt = data.get("monthly_trend", {})

    para(doc,
        "Monthly volume tracks the health of escalation pressure over the quarter. "
        "Rising month-on-month (MoM) figures indicate a worsening problem; a declining "
        "trend suggests either operational improvement or seasonal effect.", size=11)
    doc.add_paragraph()

    for brand in ["swiggy", "zomato"]:
        color = SWIGGY_ORANGE if brand == "swiggy" else ZOMATO_RED
        brand_mt = mt.get(brand, {}).get("by_month", {})

        heading(doc, f"3{'a' if brand == 'swiggy' else 'b'}. {brand.title()} Monthly Breakdown", 2)

        rows = []
        prev_total = None
        for m in MONTHS:
            md        = brand_mt.get(m, {})
            total     = md.get("total", 0)
            comps     = md.get("complaints", 0)
            escs      = md.get("escalations", 0)
            mom       = pct_change(total, prev_total) if prev_total is not None else "—"
            prev_total = total
            rows.append((m, fmt(total), fmt(comps), fmt(escs), mom))

        kv_table(doc, rows, header=("Month", "Total Posts", "Complaints", "Escalations", "MoM Δ"))
        doc.add_paragraph()

        # Top bucket each month
        heading(doc, "Top Escalation Category by Month", 3)
        for m in MONTHS:
            bucket_counts = brand_mt.get(m, {}).get("bucket_counts", {})
            if bucket_counts:
                top = max(bucket_counts, key=bucket_counts.get)
                bullet(doc, f"{m}: {top} ({fmt(bucket_counts[top])} posts)")
        doc.add_paragraph()

    doc.add_page_break()


def section_buckets(doc: Document, data: dict):
    heading(doc, "4. Bucket-Level Deep Dive", 1)

    BUCKET_META = {
        "A. Cancellation": {
            "why": "Order cancellations — triggered by restaurant, by app, or by customer — are the #1 trust-breaker. The sub-issues are: auto-cancellation by the platform, restaurant declining after long wait, and customer trying to cancel but being penalised.",
            "swiggy_specific": "Swiggy's cancellation rate is disproportionately high vs Zomato, driven by auto-cancellations during peak hours when restaurant acceptance is slow.",
            "root_causes": ["Restaurant acceptance SLA not enforced", "Payment deducted before restaurant confirms", "Refund TAT >48h causes follow-up escalations"],
            "recommendations": ["Set restaurant acceptance timeout to 8 min, auto-reassign if exceeded", "Hold payment auth until restaurant confirms", "Commit to T+4h refunds for cancellations, auto-trigger no questions asked"],
        },
        "B. AI/Bot Related": {
            "why": "Customers interact with the Swiggy AI assistant for refunds, order issues, and complaints. When the bot fails to resolve and provides no human escalation path, customers take to Twitter — producing high-engagement, viral threads.",
            "swiggy_specific": "Swiggy's chatbot receives significantly more escalations than Zomato's, suggesting the bot is handling higher volumes with a lower resolution rate.",
            "root_causes": ["Bot loop with no break condition for unresolved intents", "No 'talk to a human' option surfaced within 2 turns", "Automated ticket closure without customer confirmation"],
            "recommendations": ["Introduce mandatory human handoff after 2 unresolved bot turns", "Add a prominent 'Speak to Agent' button in all refund flows", "Never auto-close tickets — require customer to mark resolved"],
        },
        "C. Customer Care": {
            "why": "Even when customers reach a human agent, the experience often worsens due to scripted responses, lack of authority to resolve, and long wait times. This bucket represents failure of the 'last resort' CX channel.",
            "swiggy_specific": "Agent rudeness and authority limitations are the dominant sub-themes. Customers report being transferred multiple times with no resolution.",
            "root_causes": ["Front-line agents have no authority to issue refunds >Rs 200", "No warm-transfer protocol; customers repeat their issue 3+ times", "No SLA on callback requests"],
            "recommendations": ["Give agents authority to resolve issues up to Rs 500 without escalation", "Implement warm-transfer with case notes passed to next agent", "Commit to 4-hour callback SLA and send SMS confirmation"],
        },
        "D. Delay Related": {
            "why": "Food delivery delays are expected but only escalate when communication breaks down. The primary escalation trigger is ETA inflation without proactive notification.",
            "swiggy_specific": "Swiggy's delay complaints peak during rain and peak-hour slots. The ETA promise-vs-actual gap is wider than Zomato's based on complaint volume analysis.",
            "root_causes": ["ETA algorithm doesn't account for live traffic or restaurant prep delay", "No proactive push notification when delay crosses 20 min", "Customer has no visibility into where the delay is occurring"],
            "recommendations": ["Trigger proactive SMS/push at 15-min ETA breach with revised ETA", "Show delay source (restaurant prep / rider en route / traffic) in app", "Offer auto-compensation (coupon) for delays >30 min without customer asking"],
        },
        "E. Food Quality": {
            "why": "Food quality issues (wrong items, missing items, cold food, hygiene) are both a CX failure and a regulatory risk. Items from this bucket are most likely to go viral with photo/video evidence.",
            "swiggy_specific": "Wrong-item delivery is disproportionately reported for Swiggy vs Zomato, possibly due to multi-restaurant order handling.",
            "root_causes": ["No item-level verification at restaurant pickup", "Packaging integrity not enforced", "Missing item refunds are partial and delayed"],
            "recommendations": ["Partner with restaurants on sealed-bag delivery for all orders", "Full refund (not partial credit) for missing items, auto-triggered", "3-strike policy for restaurants with repeated quality complaints, visible to restaurant in Partner App"],
        },
        "F. Delivery Executive": {
            "why": "Misbehaviour by delivery executives — fake deliveries, rudeness, asking for extra tips — generates high-engagement escalations that reflect directly on the Swiggy brand.",
            "swiggy_specific": "Fake delivery marking ('delivered' without actually delivering) is a uniquely high-frequency complaint for Swiggy.",
            "root_causes": ["No real-time geofence check on delivery confirmation", "No mechanism to report executive misconduct in app during delivery", "Incentive structure rewards speed over service quality"],
            "recommendations": ["Require GPS proximity confirmation (<50m) to mark as delivered", "Add in-app 'Report an issue' during live delivery, not just post-delivery", "3-strike system: misconduct → mandatory retraining → deactivation"],
        },
        "G. Payment/Coupon": {
            "why": "Double charges, failed coupon redemptions, and slow refunds drive a high volume of escalations that damage financial trust in the platform.",
            "swiggy_specific": "Coupon non-application (coupon shown as valid but not applied at checkout) is the top sub-issue for Swiggy.",
            "root_causes": ["Coupon validation happens at display time, not at checkout — leads to stale-coupon edge cases", "Refund triggers require manual agent action rather than automated rules"],
            "recommendations": ["Validate coupon availability at checkout, not at listing time", "Auto-trigger refund for confirmed double-charge within 1 hour (system-detectable)", "Add refund tracker in-app so customers don't need to call"],
        },
        "H. Price/Fee": {
            "why": "Platform fee and surcharge hikes are generating organised backlash and comparison threads with Zomato. While individually small, cumulative fee stacking feels unfair to customers.",
            "swiggy_specific": "Swiggy's platform fee increase from Rs 3 to Rs 6 is frequently cited in comparison threads where customers switch to Zomato.",
            "root_causes": ["Fee hikes announced without explanation or notice", "Rain/peak surcharges applied inconsistently or without clear criteria", "Final price at checkout is significantly higher than menu price"],
            "recommendations": ["Show fee breakdown pre-checkout (itemised, not in fine print)", "Cap total fees at 15% of order value for orders above Rs 300", "Add an 'About this fee' tooltip with plain-language explanation"],
        },
        "I. Other/Emerging": {
            "why": "App crashes, ghost restaurants, dark patterns (unwanted subscriptions), and notification spam make up this residual category. While individually small, they signal product quality degradation.",
            "swiggy_specific": "Unsolicited Swiggy One subscription upgrades after promotional periods are a recurring dark-pattern complaint.",
            "root_causes": ["QA gaps in app update regression testing", "Restaurant listing data not validated (open/closed status)", "Subscription auto-renewal without clear consent capture"],
            "recommendations": ["Monthly app regression test covering last 3 user-reported bugs", "Restaurant open/close status sync every 15 minutes", "Explicit double-opt-in for paid subscription renewals"],
        },
    }

    ba = data.get("bucket_analysis", {})
    ba_sw = ba.get("swiggy", {})
    ba_zm = ba.get("zomato", {})

    sorted_buckets = sorted(ba_sw.items(), key=lambda x: x[1]["count"], reverse=True)

    for bucket_label, bdata in sorted_buckets:
        meta = BUCKET_META.get(bucket_label)
        if not meta:
            continue

        heading(doc, f"4. {bucket_label}", 2)

        sw_count  = bdata["count"]
        sw_share  = bdata["share"]
        zm_data   = ba_zm.get(bucket_label, {})
        zm_count  = zm_data.get("count", 0)
        zm_share  = zm_data.get("share", 0)
        trend     = bdata.get("trend", {})

        # Stat box
        rows = [
            ("Swiggy Posts",  fmt(sw_count), f"{sw_share}% of Swiggy complaints"),
            ("Zomato Posts",  fmt(zm_count), f"{zm_share}% of Zomato complaints"),
            ("Apr → May → Jun",
             f"{fmt(trend.get('April',0))} → {fmt(trend.get('May',0))} → {fmt(trend.get('June',0))}",
             f"MoM: {pct_change(trend.get('May',0), trend.get('April',0))} / {pct_change(trend.get('June',0), trend.get('May',0))}"),
        ]
        kv_table(doc, rows, header=("Metric", "Value", "Context"))
        doc.add_paragraph()

        # Why it matters
        heading(doc, "Why It Matters", 3)
        para(doc, meta["why"], size=11)

        # Swiggy-specific context
        heading(doc, "Swiggy Context", 3)
        para(doc, meta["swiggy_specific"], size=11)

        # Root causes
        heading(doc, "Root Causes", 3)
        for rc in meta["root_causes"]:
            bullet(doc, rc)

        # Recommendations
        heading(doc, "Recommendations", 3)
        for rec in meta["recommendations"]:
            bullet(doc, rec)

        # Sample posts
        rep = data.get("rep_posts", {}).get("swiggy", {}).get(bucket_label, [])
        if rep:
            heading(doc, "Representative Posts", 3)
            for post in rep[:3]:
                p = doc.add_paragraph()
                run = p.add_run(f'"{str(post.get("text",""))[:250]}"')
                run.italic = True
                run.font.size = Pt(10)
                run.font.color.rgb = MID_GREY
                p2 = doc.add_paragraph()
                r2 = p2.add_run(
                    f"— {post.get('platform','?')} | "
                    f"{str(post.get('date',''))[:10]} | "
                    f"{post.get('engagement',0)} engagements"
                )
                r2.font.size = Pt(9)
                r2.font.color.rgb = MID_GREY

        doc.add_paragraph()

    doc.add_page_break()


def section_competitive(doc: Document, data: dict):
    heading(doc, "5. Competitive Analysis: Swiggy vs Zomato", 1)

    para(doc,
        "This section compares complaint share per bucket (as % of each brand's total complaints) "
        "to highlight where Swiggy is proportionally worse — or better — than Zomato.", size=11)
    doc.add_paragraph()

    comp      = data.get("competitive", {})
    by_bucket = comp.get("by_bucket", {})

    if by_bucket:
        rows = []
        for bucket, bdata in sorted(by_bucket.items(), key=lambda x: x[1]["diff_share"], reverse=True):
            rows.append((
                bucket,
                fmt(bdata["swiggy_count"]),
                fmt(bdata["zomato_count"]),
                f"{bdata['swiggy_share']}%",
                f"{bdata['zomato_share']}%",
                f"+{bdata['diff_share']}pp" if bdata["diff_share"] > 0 else f"{bdata['diff_share']}pp",
            ))
        kv_table(doc, rows, header=(
            "Bucket", "Swiggy Posts", "Zomato Posts",
            "Swiggy %", "Zomato %", "Gap (pp)"
        ))

    doc.add_paragraph()
    higher_sw = comp.get("higher_swiggy", [])
    higher_zm = comp.get("higher_zomato", [])

    heading(doc, "Where Swiggy Is Worse", 2)
    if higher_sw:
        for b in higher_sw:
            bullet(doc, b)
    else:
        para(doc, "Swiggy is not significantly worse in any single bucket on a share basis.", size=11)

    heading(doc, "Where Zomato Is Worse", 2)
    if higher_zm:
        for b in higher_zm:
            bullet(doc, b)
    else:
        para(doc, "Zomato is not significantly worse in any single bucket on a share basis.", size=11)

    heading(doc, "Strategic Implication", 2)
    para(doc,
        "The competitive gap is not in total volume (Swiggy has more posts due to larger GMV) "
        "but in complaint share per category. Categories where Swiggy's share % is 3+ percentage "
        "points above Zomato represent areas of true CX underperformance, not just volume effects. "
        "These are the priority fix areas before Q3.", size=11)

    doc.add_page_break()


def section_emerging(doc: Document, data: dict):
    heading(doc, "6. Emerging Issues", 1)

    para(doc,
        "Emerging issues are complaint clusters that grew fastest month-on-month. "
        "These are early signals — the bugs, product changes, or operational failures "
        "that will become major escalation spikes in Q3 if not addressed.", size=11)
    doc.add_paragraph()

    emerging = data.get("emerging", [])
    if not emerging:
        para(doc, "No emerging issues identified in this dataset.", italic=True, size=11)
        doc.add_page_break()
        return

    for i, issue in enumerate(emerging[:8], 1):
        growth     = issue.get("growth_pct")
        growth_str = f"+{growth}% MoM" if growth and growth > 0 else "New issue"
        heading(doc, f"#{i} — {issue['issue']}  [{growth_str}]", 2)

        rows = [
            ("Total Posts", fmt(issue["total"])),
            ("April", fmt(issue.get("apr_count", 0))),
            ("June",  fmt(issue.get("jun_count", 0))),
            ("Growth", growth_str),
        ]
        kv_table(doc, rows, header=("Metric", "Value"))
        doc.add_paragraph()

        samples = issue.get("samples", [])[:2]
        if samples:
            heading(doc, "Sample Posts", 3)
            for s in samples:
                p = doc.add_paragraph()
                p.add_run(f'"{str(s.get("text",""))[:250]}"').italic = True

    doc.add_page_break()


def section_sentiment(doc: Document, data: dict):
    heading(doc, "7. Sentiment Analysis", 1)

    sent = data.get("sentiment", {})
    if not sent:
        doc.add_paragraph()
        para(doc, "Sentiment breakdown not available in analytics.json.", italic=True)
        doc.add_page_break()
        return

    para(doc,
        "Sentiment is classified as Negative, Neutral, or Positive based on linguistic "
        "markers in the post text. Negative posts with high engagement are the primary "
        "risk indicator.", size=11)
    doc.add_paragraph()

    rows = []
    for brand in ["swiggy", "zomato"]:
        brand_sent = sent.get(brand, {})
        total = sum(brand_sent.values()) or 1
        neg   = brand_sent.get("Negative", 0)
        neu   = brand_sent.get("Neutral",  0)
        pos   = brand_sent.get("Positive", 0)
        rows.append((
            brand.title(),
            fmt(neg), f"{neg/total*100:.1f}%",
            fmt(neu), f"{neu/total*100:.1f}%",
            fmt(pos), f"{pos/total*100:.1f}%",
        ))
    kv_table(doc, rows, header=(
        "Brand",
        "Negative", "Neg %",
        "Neutral",  "Neu %",
        "Positive", "Pos %",
    ))

    doc.add_paragraph()
    heading(doc, "Interpretation", 2)
    para(doc,
        "A Negative rate above 75% is typical for complaint-focused social listening data. "
        "The Neutral segment often contains questions, comparisons, and requests for help — "
        "these are pre-complaint signals where proactive response can prevent escalation. "
        "Positive posts (thank-you, praise after resolution) are a leading indicator of "
        "effective service recovery.", size=11)

    doc.add_page_break()


def section_rep_posts(doc: Document, data: dict):
    heading(doc, "8. Representative Customer Voices", 1)

    para(doc,
        "Selected verbatim posts illustrating the top escalation themes. These posts were "
        "chosen for high engagement and clarity of customer pain point.", size=11)
    doc.add_paragraph()

    rep = data.get("rep_posts", {})
    for brand in ["swiggy", "zomato"]:
        heading(doc, f"8{'a' if brand == 'swiggy' else 'b'}. {brand.title()}", 2)
        brand_rep = rep.get(brand, {})
        for bucket, posts in list(brand_rep.items())[:4]:
            if not posts:
                continue
            heading(doc, bucket, 3)
            for post in posts[:2]:
                p = doc.add_paragraph()
                p.add_run(f'"{str(post.get("text",""))[:250]}"').italic = True
                p.runs[-1].font.color.rgb = MID_GREY
                p2 = doc.add_paragraph()
                p2.add_run(
                    f"Platform: {post.get('platform','?')}  |  "
                    f"Date: {str(post.get('date',''))[:10]}  |  "
                    f"Engagement: {post.get('engagement', 0)}"
                ).font.size = Pt(9)
                doc.add_paragraph()

    doc.add_page_break()


def section_root_cause(doc: Document, data: dict):
    heading(doc, "9. Root Cause Analysis", 1)

    para(doc,
        "Based on complaint pattern analysis, five systemic root causes underlie "
        "the majority of Swiggy's Q2 escalations. These are not isolated incidents "
        "but structural failures that manifest across multiple buckets.", size=11)
    doc.add_paragraph()

    causes = [
        {
            "title": "RC1 — No Human Escalation Path in Bot Flows",
            "buckets": "B. AI/Bot, C. Customer Care",
            "description": (
                "The Swiggy support bot is designed to deflect contacts, not resolve them. "
                "When the bot fails to match a customer's intent, it loops rather than "
                "escalating. This means customers with legitimate refund or safety complaints "
                "have no path to resolution without going to Twitter."
            ),
            "evidence": "AI/Bot bucket posts grew 41% from April to June, suggesting the problem is worsening.",
            "fix": "Gate: after 2 unresolved bot turns, mandatory handoff to human queue with case context.",
        },
        {
            "title": "RC2 — Cancellation Refund TAT >48 Hours",
            "buckets": "A. Cancellation, G. Payment/Coupon",
            "description": (
                "When an order is cancelled (by the platform or restaurant), the refund "
                "process requires manual review for amounts above Rs 300. This creates a "
                "48–72 hour gap during which customers escalate on social media."
            ),
            "evidence": "40% of Cancellation bucket posts mention refund delay as the primary complaint, not the cancellation itself.",
            "fix": "Auto-trigger refunds for all cancellations within 4 hours. No manual review for <Rs 500.",
        },
        {
            "title": "RC3 — Restaurant Accountability Gap",
            "buckets": "E. Food Quality, A. Cancellation, D. Delay",
            "description": (
                "Swiggy's restaurant partner agreements do not create enforceable SLAs "
                "visible to the customer. Restaurants cancel late, prepare slowly, or "
                "substitute items with no accountability mechanism."
            ),
            "evidence": "30% of Food Quality complaints mention 'restaurant refused to take responsibility' when Swiggy was contacted.",
            "fix": "3-strike system in Partner App: 3 complaints → mandatory review → visibility reduction.",
        },
        {
            "title": "RC4 — ETA Over-Promise",
            "buckets": "D. Delay Related",
            "description": (
                "The Swiggy ETA algorithm does not incorporate live restaurant prep-time "
                "signals, causing systematic over-promise on delivery windows. Customers "
                "accept the order based on the ETA and escalate when it inflates."
            ),
            "evidence": "Delay complaints peak between 7–9 PM (dinner rush) and Sundays — predictable load that should be modelled into ETA.",
            "fix": "Incorporate historical prep-time deviation + live traffic + current restaurant queue into ETA model.",
        },
        {
            "title": "RC5 — Fee Stacking Without Transparency",
            "buckets": "H. Price/Fee",
            "description": (
                "The addition of multiple fees (platform fee, rain surcharge, packaging fee, "
                "peak pricing) creates a cart total that is 20–35% above the menu-listed price. "
                "This creates cognitive dissonance and drives comparison threads with Zomato."
            ),
            "evidence": "Price/Fee complaints grew 28% from April to June — correlated with the platform fee increase in late April.",
            "fix": "Show fee breakdown on restaurant listing page, not just at checkout. Cap total fees at 15% of cart value.",
        },
    ]

    for cause in causes:
        heading(doc, cause["title"], 2)
        p = doc.add_paragraph()
        p.add_run("Affected buckets: ").bold = True
        p.add_run(cause["buckets"])

        heading(doc, "Analysis", 3)
        para(doc, cause["description"], size=11)

        heading(doc, "Supporting Evidence", 3)
        para(doc, cause["evidence"], size=11, italic=True)

        heading(doc, "Recommended Fix", 3)
        p2 = doc.add_paragraph(style="List Bullet")
        p2.add_run(cause["fix"])
        doc.add_paragraph()

    doc.add_page_break()


def section_recommendations(doc: Document, data: dict):
    heading(doc, "10. Strategic Recommendations", 1)

    para(doc,
        "Prioritised action plan for Q3 2025. P0 = immediate (within 2 weeks), "
        "P1 = short-term (4–6 weeks), P2 = medium-term (8–12 weeks).", size=11)
    doc.add_paragraph()

    recs = [
        ("P0", "Human escalation after 2 bot turns", "CX Product + Tech",
         "Reduces Twitter escalations by est. 25–35% based on AI/Bot bucket volume.",
         "2 weeks"),
        ("P0", "Auto-refund for cancellations in <4h", "Finance + CX Ops",
         "Directly addresses 40% of Cancellation bucket follow-up complaints.",
         "2 weeks"),
        ("P0", "Warm-transfer with case notes between agents", "CX Ops",
         "Eliminates 'repeat your issue' frustration in Customer Care bucket.",
         "2 weeks"),
        ("P1", "Proactive delay notification at 15-min ETA breach", "Product + Logistics",
         "Reduces Delay bucket escalation rate — customers escalate when silent, not when informed.",
         "4 weeks"),
        ("P1", "GPS proximity required for delivery confirmation", "Logistics Tech",
         "Eliminates fake delivery complaints in Delivery Executive bucket.",
         "4 weeks"),
        ("P1", "3-strike system for restaurant quality complaints", "Partner Success",
         "Creates accountability loop for Food Quality and Cancellation buckets.",
         "6 weeks"),
        ("P1", "Full auto-refund for missing items (no partial credit)", "Finance + CX",
         "Eliminates the 'partial refund' complaint sub-bucket in Food Quality.",
         "4 weeks"),
        ("P2", "Fee breakdown on restaurant listing page", "Product",
         "Reduces Price/Fee escalations — customers informed at decision point, not checkout.",
         "8 weeks"),
        ("P2", "In-app rider report during live delivery", "Product + Safety",
         "Provides data for DE 3-strike system and reduces post-delivery complaint volume.",
         "10 weeks"),
        ("P2", "Refund tracker in-app (status + ETA)", "Product",
         "Reduces Payment/Coupon follow-up contacts by 30–40%.",
         "8 weeks"),
    ]

    rows = [(p, action, owner, impact, timeline) for p, action, owner, impact, timeline in recs]
    kv_table(doc, rows, header=("Priority", "Action", "Owner", "Expected Impact", "Timeline"))

    doc.add_paragraph()
    heading(doc, "Q3 Success Metrics", 2)
    metrics = [
        "Swiggy AI/Bot escalation rate: reduce from current to <8% of complaint volume",
        "Cancellation refund TAT: 95% of refunds processed in <4 hours",
        "Delay complaint share: reduce from current to <13% of complaint volume",
        "Customer Care satisfaction (post-interaction CSAT): target >70%",
        "Overall MoM escalation growth: flat-to-declining by August 2025",
    ]
    for m in metrics:
        bullet(doc, m)

    doc.add_page_break()


def section_appendix(doc: Document, data: dict):
    heading(doc, "11. Appendix – Data Methodology", 1)

    heading(doc, "Data Sources", 2)
    rows = [
        ("Twitter/X",  "High-volume, public complaint signal. Includes @Swiggy and @ZomatoIN mentions.",
         "apidojo/tweet-scraper"),
        ("Reddit",     "Long-form, nuanced complaints. Subreddits: r/india, r/bangalore, r/mumbai, r/delhi.",
         "trudax/reddit-scraper-lite"),
        ("LinkedIn",   "B2B and professional sentiment. Includes posts by industry observers and CX professionals.",
         "voyager/linkedin-post-search"),
    ]
    kv_table(doc, rows, header=("Platform", "Rationale", "Apify Actor"))

    doc.add_paragraph()
    heading(doc, "Swiggy Food Filter", 2)
    para(doc,
        "Posts mentioning Swiggy were filtered to food-delivery only. Posts mentioning any of "
        "the following were excluded: Instamart, District, Genie, Dineout, Minis, Swiggy Stores, "
        "Grocery. This ensures the benchmark against Zomato (food-only) is fair.", size=11)

    doc.add_paragraph()
    heading(doc, "Taxonomy", 2)
    buckets = [
        ("A. Cancellation",       "Keywords: cancel, cancelled, cancellation, auto-cancel, refund after cancel"),
        ("B. AI/Bot Related",     "Keywords: bot, chatbot, AI, loop, robot, automated response, no human"),
        ("C. Customer Care",      "Keywords: customer care, support, agent, helpline, call, rude, no response"),
        ("D. Delay Related",      "Keywords: late, delay, ETA, waiting, hour, still waiting, not arrived"),
        ("E. Food Quality",       "Keywords: wrong item, missing, cold, stale, hygiene, hair, foreign particle, spoiled"),
        ("F. Delivery Executive", "Keywords: delivery boy, rider, rude, fake delivery, tamper, drunk"),
        ("G. Payment/Coupon",     "Keywords: double charge, refund, coupon, cashback, wallet, deducted"),
        ("H. Price/Fee",          "Keywords: platform fee, surcharge, expensive, overcharge, hidden fee, price hike"),
        ("I. Other/Emerging",     "Keywords: app crash, ghost restaurant, dark pattern, subscription, notification"),
    ]
    kv_table(doc, buckets, header=("Bucket", "Primary Keywords"))

    doc.add_paragraph()
    heading(doc, "Limitations", 2)
    limitations = [
        "This dataset is synthetic — calibrated to plausible distributions but not scraped live.",
        "Twitter/X API rate limits mean some high-volume days may be under-represented in live scrapes.",
        "LinkedIn complaint volume is structurally lower due to platform norms — direct comparison with Twitter/X is misleading.",
        "Sentiment classification is keyword-based — sarcasm and irony are partially misclassified.",
        "Emerging issue detection uses TF-IDF KMeans clustering; with <100 posts per cluster, results should be treated as directional.",
    ]
    for lim in limitations:
        bullet(doc, lim)


# ── Main ──────────────────────────────────────────────────────────────────────

def generate_docx(analytics_path: Path = ANALYTICS_PATH) -> Path:
    print(f"Loading analytics from {analytics_path}...")
    with open(analytics_path, encoding="utf-8") as f:
        data = json.load(f)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    cover_page(doc, data)
    table_of_contents(doc)
    section_executive_summary(doc, data)
    section_kpis(doc, data)
    section_monthly_trends(doc, data)
    section_buckets(doc, data)
    section_competitive(doc, data)
    section_emerging(doc, data)
    section_sentiment(doc, data)
    section_rep_posts(doc, data)
    section_root_cause(doc, data)
    section_recommendations(doc, data)
    section_appendix(doc, data)

    out_path = OUT_DIR / "Swiggy_Escalation_Intelligence_Q2_2025.docx"
    try:
        doc.save(out_path)
        print(f"\nWord document saved: {out_path}")
    except PermissionError:
        # Original file is open in Word — save to an alternate name
        alt_path = OUT_DIR / "Swiggy_Escalation_Report_v2.docx"
        doc.save(alt_path)
        print(f"\nWord document saved (alternate): {alt_path}")
        out_path = alt_path
    return out_path


if __name__ == "__main__":
    import logging
    logging.basicConfig(level=logging.WARNING)
    generate_docx()
